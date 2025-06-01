import os
from dotenv import load_dotenv
load_dotenv()  # Carrega variáveis do .env

import json
import base64
from datetime import datetime, timedelta, timezone
import tempfile
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file, Response
import firebase_admin
from firebase_admin import credentials, auth, db, storage
import pyrebase
from functools import wraps
import uuid
import requests
import logging
from groq import Groq
from werkzeug.utils import secure_filename
from markupsafe import Markup
import markdown
import stripe
import subprocess
import platform



# Configuração do Flask
app = Flask(__name__)
app.secret_key = os.getenv('FLASK_SECRET_KEY')
app.config['PERMANENT_SESSION_LIFETIME'] = timedelta(days=30)  # Sessão dura 30 dias

# Configuração de logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)


import json
import base64
from datetime import datetime, timedelta
import tempfile
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, send_file, Response
import firebase_admin
from firebase_admin import credentials, auth, db, storage
import pyrebase
from functools import wraps
import uuid
import requests
import logging
from groq import Groq
from werkzeug.utils import secure_filename
from markupsafe import Markup
import markdown
import stripe
from datetime import datetime
from flask import send_file
from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE
from flask import jsonify, render_template_string

# Lista de campos importantes que devem ser destacados
IMPORTANT_FIELDS = [
    'Público-Alvo:', 'Introdução:', 'Objetivo:', 'Descrição:', 'Resultado:',
    'Status:', 'Tipo:', 'Template:', 'Data:', 'Imagem de Referência:',
    'Campos Adicionais:', 'Descrição do Projeto:', 'Resultado do Briefing:',
    'Informações do Briefing:', 'Campos Adicionais:'
]

# Função auxiliar para verificar se um texto é um campo importante
def is_important_field(text):
    return any(text.strip().startswith(field) for field in IMPORTANT_FIELDS)

# Inicialização do Pyrebase para autenticação de usuário final
firebase_config = {
    "apiKey": os.getenv("FIREBASE_API_KEY"),
    "authDomain": os.getenv("FIREBASE_AUTH_DOMAIN"),
    "databaseURL": os.getenv("FIREBASE_DATABASE_URL"),
    "projectId": os.getenv("FIREBASE_PROJECT_ID"),
    "storageBucket": os.getenv("FIREBASE_STORAGE_BUCKET"),
    "messagingSenderId": os.getenv("FIREBASE_MESSAGING_SENDER_ID"),
    "appId": os.getenv("FIREBASE_APP_ID"),
}
pyrebase_app = pyrebase.initialize_app(firebase_config)
pyre_auth = pyrebase_app.auth()

# Configuração dos planos e limites
PLANS = {
    'free': {
        'name': 'Grátis',
        'price': 0,
        'brief_limit': 3,
        'features': [
            '3 briefs por mês',
            'Templates básicos',
            'Exportação em PDF',
            'Suporte por email'
        ],
        'templates': ['classic'],
        'export_formats': ['pdf', 'docx'],
        'ai_models': ['gemini']
    },
    'pro': {
        'name': 'Pro',
        'price': 9,
        'stripe_price_id': os.getenv('STRIPE_PRO_PRICE_ID'),
        'brief_limit': 50,
        'features': [
            '50 briefs por mês',
            'Templates avançados',
            'Exportação em múltiplos formatos',
            'Suporte prioritário',
            'Dashboard avançado',
            'Histórico completo',
            'Personalização de templates',
            'Modelos de IA avançados',
            'Campos extras personalizados'
        ],
        'templates': ['classic', 'visual', 'minimal'],
        'export_formats': ['pdf', 'docx', 'txt', 'html'],
        'ai_models': ['gemini', 'groq']
    },
    'premium': {
        'name': 'Premium',
        'price': 19,
        'stripe_price_id': os.getenv('STRIPE_PREMIUM_PRICE_ID'),
        'brief_limit': 50,  # Mantendo o limite de 50 para manter consistência com o plano Pro
        'features': [
            '50 briefs por mês',
            'Templates exclusivos',
            'Exportação em massa',
            'Suporte VIP',
            'Dashboard completo',
            'Backup automático',
            'API access',
            'Equipes e colaboração',
            'Personalização avançada',
            'Todos os modelos de IA'
        ],
        'templates': ['classic', 'visual', 'minimal', 'exclusive'],
        'export_formats': ['pdf', 'docx', 'txt', 'html'],
        'ai_models': ['gemini', 'groq']
    }
}

stripe.api_key = os.getenv('STRIPE_SECRET_KEY')
stripe_publishable_key = os.getenv('STRIPE_PUBLISHABLE_KEY')
if not stripe.api_key or not stripe_publishable_key:
    print("Warning: Stripe keys not found in environment variables")

# Função auxiliar para obter plano do usuário e briefs usados
from datetime import datetime

def get_user_plan_info(user_id):
    user_ref = db.reference(f'users/{user_id}')
    user_data = user_ref.get() or {}
    plan = user_data.get('plan', 'free')
    plan_expiry = user_data.get('plan_expiry')
    # Reset mensal do contador de briefs
    now = datetime.utcnow()
    month_key = now.strftime('%Y-%m')
    briefs_used = user_data.get('briefs_used', {})
    briefs_this_month = briefs_used.get(month_key, 0)
    return plan, plan_expiry, briefs_this_month, user_data

def increment_brief_count(user_id):
    user_ref = db.reference(f'users/{user_id}')
    user_data = user_ref.get() or {}
    now = datetime.utcnow()
    month_key = now.strftime('%Y-%m')
    briefs_used = user_data.get('briefs_used', {})
    briefs_used[month_key] = briefs_used.get(month_key, 0) + 1
    user_ref.update({'briefs_used': briefs_used})

# Função para checar se usuário pode criar brief

def can_create_brief(user_id):
    plan, plan_expiry, briefs_this_month, user_data = get_user_plan_info(user_id)
    limit = PLANS[plan]['brief_limit']
    
    # Verifica se o plano expirou e faz downgrade para free
    if plan_expiry:
        expiry_dt = datetime.strptime(plan_expiry, '%Y-%m-%d')
        if expiry_dt < datetime.utcnow():
            # Downgrade automático
            db.reference(f'users/{user_id}').update({
                'plan': 'free',
                'plan_expiry': None,
                'briefs_used': {},  # Reseta o contador de briefs
                'notified_limit': {}  # Reseta os avisos de limite
            })
            plan = 'free'
            limit = PLANS['free']['brief_limit']
            
            # Notifica sobre o downgrade
            email = user_data.get('email')
            if email:
                try:
                    send_email(
                        subject='Aviso: Plano expirado - EazyBrief',
                        recipients=[email],
                        body=f'Seu plano {PLANS[user_data.get("plan", "free")]["name"]} expirou. Você foi automaticamente revertido para o plano gratuito. Faça upgrade para continuar usando sem restrições.'
                    )
                except Exception as e:
                    print(f'Erro ao enviar email de downgrade: {e}')
    
    # Verifica se atingiu o limite de briefs
    if limit is not None:  # Se tiver limite definido (free e pro)
        now = datetime.utcnow()
        month_key = now.strftime('%Y-%m')
        if briefs_this_month >= limit:
            user_ref = db.reference(f'users/{user_id}')
            user_data = user_ref.get() or {}
            email = user_data.get('email')
            notified = user_data.get('notified_limit', {})
            if not notified.get(month_key):
                if email:
                    try:
                        send_email(
                            subject='Atenção: Limite do plano atingido - EazyBrief',
                            recipients=[email],
                            body=f'Você atingiu o limite de briefs do seu plano {PLANS[plan]["name"]}. Faça upgrade para continuar usando sem restrições.'
                        )
                    except Exception as e:
                        print(f'Erro ao enviar email de limite atingido: {e}')
                notified[month_key] = True
                user_ref.update({'notified_limit': notified})
            return False
    
    return True


# Configuração do Firebase
firebase_config = {
    "apiKey": os.getenv('FIREBASE_API_KEY'),
    "authDomain": os.getenv('FIREBASE_AUTH_DOMAIN'),
    "projectId": os.getenv('FIREBASE_PROJECT_ID'),
    "storageBucket": os.getenv('FIREBASE_STORAGE_BUCKET'),
    "messagingSenderId": os.getenv('FIREBASE_MESSAGING_SENDER_ID'),
    "appId": os.getenv('FIREBASE_APP_ID'),
    "databaseURL": os.getenv('FIREBASE_DATABASE_URL')
}

# Inicialização do Firebase Admin
cred_json = os.getenv('FIREBASE_CREDENTIALS_JSON')
if not cred_json:
    raise ValueError("FIREBASE_CREDENTIALS_JSON não configurado nas variáveis de ambiente")

try:
    cred_dict = json.loads(cred_json)
    cred = credentials.Certificate(cred_dict)
    firebase_admin.initialize_app(cred, {
        'databaseURL': firebase_config['databaseURL'],
        'storageBucket': os.getenv('FIREBASE_STORAGE_BUCKET')
    })
except json.JSONDecodeError:
    raise ValueError("FIREBASE_CREDENTIALS_JSON contém JSON inválido")
except Exception as e:
    raise ValueError(f"Erro ao inicializar Firebase Admin: {str(e)}")

# Configuração do Groq
api_key = os.getenv('GROQ_API_KEY')
print(f"API Key carregada: {'Sim' if api_key else 'Não'}")  # Debug
print(f"API Key: {api_key[:10]}...")  # Mostra apenas os primeiros 10 caracteres por segurança

groq_client = Groq(api_key=api_key)

# Filtro markdown
@app.template_filter('markdown')
def markdown_filter(text):
    return Markup(markdown.markdown(text or '', extensions=['extra', 'smarty']))

# Função auxiliar para obter user_id
def get_user_id():
    if 'user' not in session or 'user_email' not in session:
        return None
    user_id = session['user'].get('uid') or session['user'].get('localId')
    if not user_id:
        user_id = session['user_email'].replace('.', '_')
    return user_id

# Decorator para rotas que requerem autenticação
def login_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' not in session or 'user_email' not in session:
            flash('Faça login para acessar esta página.', 'warning')
            return redirect(url_for('login'))
        return f(*args, **kwargs)
    return decorated_function

# Decorator para rotas que requerem usuário não autenticado
def guest_required(f):
    @wraps(f)
    def decorated_function(*args, **kwargs):
        if 'user' in session and 'user_email' in session:
            flash('Você já está logado.', 'info')
            return redirect(url_for('dashboard'))
        return f(*args, **kwargs)
    return decorated_function

@app.route('/')
def index():
    if 'user' in session and 'user_email' in session:
        return redirect(url_for('dashboard'))
    return render_template('index.html')

@app.route('/signup', methods=['GET', 'POST'])
@guest_required
def signup():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        if not email or not password:
            flash('Email e senha são obrigatórios.', 'danger')
            return render_template('signup.html')
        try:
            # Criar usuário usando Pyrebase
            user = pyre_auth.create_user_with_email_and_password(email, password)
            
            # Enviar email de verificação usando Pyrebase
            pyre_auth.send_email_verification(user['idToken'])
            
            flash('Conta criada com sucesso! Verifique seu email para ativar sua conta.', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            error_message = str(e)
            print(f"Erro detalhado no signup: {error_message}")
            if 'EMAIL_EXISTS' in error_message:
                flash('Este email já está em uso.', 'danger')
            else:
                flash('Erro ao criar conta. Por favor, tente novamente.', 'danger')
    return render_template('signup.html')

@app.route('/login', methods=['GET', 'POST'])
@guest_required
def login():
    if 'email_verified' in session and session['email_verified']:
        flash('Email verificado com sucesso!', 'success')
        session.pop('email_verified', None)
        session.pop('verified_email', None)
    
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        password = request.form.get('password', '').strip()
        resend_verification = request.form.get('resend_verification', '').lower() == 'true'
        
        try:
            if resend_verification:
                # Gerar novo link de verificação
                action_code_settings = auth.ActionCodeSettings(
                    url=f"{request.url_root}auth_action?mode=verifyEmail",
                    handle_code_in_app=True
                )
                auth.generate_email_verification_link(email, action_code_settings=action_code_settings)
                flash('Email de verificação reenviado. Por favor, verifique sua caixa de entrada.', 'info')
                return render_template('login.html', show_resend_button=True, email=email)
            
            # Tentar login
            user = pyre_auth.sign_in_with_email_and_password(email, password)
            user_info = auth.get_user_by_email(email)
            
            if not user_info.email_verified:
                flash('Verifique seu email antes de acessar o dashboard.', 'warning')
                return render_template('login.html', show_resend_button=True, email=email)
            
            session.permanent = True  # Torna a sessão permanente
            session['user'] = {'uid': user_info.uid, 'refreshToken': user['refreshToken']}
            session['user_email'] = email
            flash('Login bem-sucedido!', 'success')
            return redirect(url_for('dashboard'))
        except Exception as e:
            flash('Email ou senha incorretos ou erro ao fazer login.', 'danger')
    return render_template('login.html', show_resend_button=False)

@app.route('/google/callback', methods=['POST'])
def google_callback():
    try:

        
        id_token = request.json.get('idToken')
        if not id_token:
            return jsonify({'success': False, 'message': 'Token não fornecido'})
        
        # Verificar o token com o Firebase Admin
        decoded_token = auth.verify_id_token(id_token)
        user_id = decoded_token['uid']
        email = decoded_token['email']
        
        # Verificar se o usuário já existe
        try:
            user = auth.get_user(user_id)
        except auth.UserNotFoundError:
            # Criar usuário se não existir
            user = auth.create_user(
                uid=user_id,
                email=email,
                email_verified=True
            )
        
        # Atualizar sessão
        session['user'] = {'uid': user_id}
        session['user_email'] = email
        
        return jsonify({'success': True})
    except Exception as e:
        logging.error(f'Erro no callback do Google: {str(e)}')
        return jsonify({'success': False, 'message': str(e)})

@app.route('/google/login')
def google_login():
    return redirect(url_for('login'))

@app.route('/reset-password', methods=['GET', 'POST'])
def reset_password():
    if request.method == 'POST':
        email = request.form.get('email', '').strip()
        if not email:
            flash('Email é obrigatório.', 'danger')
            return render_template('reset_password.html')
        
        try:
            # Gerar link de redefinição de senha
            action_code_settings = auth.ActionCodeSettings(
                url=f"{request.url_root}auth_action?mode=resetPassword",
                handle_code_in_app=True
            )
            auth.generate_password_reset_link(email, action_code_settings=action_code_settings)
            flash('Link de redefinição de senha enviado. Por favor, verifique sua caixa de entrada.', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            print(f'Erro ao enviar email de redefinição: {str(e)}')
            flash('Erro ao enviar email de redefinição. Por favor, tente novamente.', 'danger')
    
    return render_template('reset_password.html')

@app.route('/auth_action')
def handle_firebase_action():
    mode = request.args.get('mode', '')
    oob_code = request.args.get('oobCode', '')
    
    print(f"DEBUG - Mode: {mode}, OOB Code: {oob_code}")
    
    if mode == 'verifyEmail' and oob_code:
        try:
            # Tentar fazer login com o código OOB
            user = pyre_auth.sign_in_with_email_link(oob_code)
            if user and user.get('email'):
                session['email_verified'] = True
                session['verified_email'] = user['email']
                flash('Email verificado com sucesso!', 'success')
                return redirect(url_for('login'))
            else:
                raise Exception("Falha ao verificar email")
        except Exception as e:
            print(f'Erro detalhado ao verificar email: {str(e)}')
            flash('Erro ao verificar email. Por favor, tente solicitar um novo email de verificação na página de login.', 'danger')
            return redirect(url_for('login'))
    elif mode == 'resetPassword' and oob_code:
        try:
            session['reset_password_code'] = oob_code
            return redirect(url_for('reset_password_confirm'))
        except Exception as e:
            flash(f'Erro ao processar link de redefinição: {str(e)}', 'danger')
            return redirect(url_for('login'))
    
    flash('Ação inválida.', 'danger')
    return redirect(url_for('login'))

@app.route('/reset-password-confirm', methods=['GET', 'POST'])
def reset_password_confirm():
    # Obter o código da sessão ou dos parâmetros da URL
    oob_code = session.get('reset_password_code') or request.args.get('oobCode')
    
    if not oob_code:
        flash('Código de redefinição de senha inválido ou expirado.', 'danger')
        return redirect(url_for('login'))
    
    if request.method == 'POST':
        new_password = request.form.get('new_password', '').strip()
        if not new_password:
            flash('Nova senha é obrigatória.', 'danger')
            return render_template('reset_password_confirm.html')
        
        try:
            # Confirmar a redefinição de senha
            pyre_auth.confirm_password_reset(oob_code, new_password)
            session.pop('reset_password_code', None)
            flash('Senha redefinida com sucesso!', 'success')
            return redirect(url_for('login'))
        except Exception as e:
            print(f'Erro ao redefinir senha: {str(e)}')
            flash('Erro ao redefinir senha. O link pode ter expirado.', 'danger')
            return redirect(url_for('login'))
    
    return render_template('reset_password_confirm.html')

@app.route('/dashboard')
@login_required
def dashboard():
    try:
        user_info = auth.get_user_by_email(session['user_email'])
        user_id = get_user_id()
        briefs = []
        briefs_count = 0
        if user_id:
            briefs_snap = db.reference(f'briefs/{user_id}').get()
            if briefs_snap:
                for key, val in briefs_snap.items():
                    val['id'] = key
                    briefs.append(val)
                briefs.sort(key=lambda b: b.get('created_at', ''), reverse=True)
                briefs_count = len(briefs)
        
        concluidos = sum(1 for b in briefs if b.get('status', 'concluido') == 'concluido')
        andamento = sum(1 for b in briefs if b.get('status', 'concluido') == 'andamento')
        
        # Criar dicionário de estatísticas
        stats = {
            'total_briefs': briefs_count,
            'completed_briefs': concluidos,
            'in_progress_briefs': andamento
        }

        # Buscar atividades recentes (últimos 2 briefs)
        recent_activity = []
        for brief in briefs[:2]:  # Pegar apenas os 2 mais recentes
            activity = {
                'description': f"Brief '{brief.get('brief_type', 'Sem tipo')}' {'concluído' if brief.get('status') == 'concluido' else 'em andamento'}",
                'timestamp': brief.get('created_at', '').split('T')[0],  # Formatar data
                'brief_id': brief['id']  # Adicionar o ID do brief
            }
            recent_activity.append(activity)
        
        return render_template('dashboard.html', 
                             user=user_info, 
                             briefs=briefs, 
                             briefs_count=briefs_count, 
                             concluidos=concluidos, 
                             andamento=andamento, 
                             stats=stats,
                             recent_activity=recent_activity)
    except Exception as e:
        print('DEBUG ERRO DASHBOARD:', repr(e))
        flash('Sessão expirada. Faça login novamente.', 'warning')
        session.clear()
        return redirect(url_for('login'))

@app.route('/logout')
def logout():
    session.clear()
    flash('Você saiu da sua conta.', 'info')
    return redirect(url_for('login'))



@app.route('/my_briefs')
@login_required
def my_briefs():
    user_info = auth.get_user_by_email(session['user_email'])
    user_id = get_user_id()
    briefs = []
    try:
        briefs_snap = db.reference(f'briefs/{user_id}').get()
        if briefs_snap:
            for key, val in briefs_snap.items():
                val['id'] = key
                briefs.append(val)
            briefs.sort(key=lambda b: b.get('created_at', ''), reverse=True)
    except Exception as e:
        print(f"Erro ao buscar briefs: {str(e)}")
    return render_template('my_briefs.html', user=user_info, briefs=briefs)

@app.route('/profile', methods=['GET', 'POST'])
@login_required
def profile():
    user_info = auth.get_user_by_email(session['user_email'])
    user_id = get_user_id()
    prefs = db.reference(f'user_prefs/{user_id}').get() or {}
    
    # Buscar informações do plano diretamente do banco de dados
    user_data = db.reference(f'users/{user_id}').get() or {}
    plan = user_data.get('plan', 'free')
    plan_expiry = user_data.get('plan_expiry')
    
    # Log para debug
    logging.info(f'[Profile] Carregando perfil para user_id={user_id}')
    logging.info(f'[Profile] Dados do usuário: {user_data}')
    logging.info(f'[Profile] Plano atual: {plan}, Expira em: {plan_expiry}')
    
    if request.method == 'POST':
        display_name = request.form.get('display_name', user_info.display_name or '')
        phone = request.form.get('phone', '').strip()
        photo_url = prefs.get('photo_url', user_info.photo_url)
        
        if 'photo' in request.files and request.files['photo'].filename:
            try:
                photo = request.files['photo']
                if photo.mimetype not in ['image/jpeg', 'image/png']:
                    flash('Formato de imagem inválido.', 'danger')
                    return redirect(url_for('profile'))
                
                filename = secure_filename(photo.filename)
                bucket_name = 'brief-generator-5c33f.firebasestorage.app'
                bucket = storage.bucket(bucket_name)
                blob = bucket.blob(f'profile_photos/{user_id}/{filename}')
                
                # Configurar metadados para tornar o arquivo público
                blob.metadata = {'cacheControl': 'public, max-age=31536000'}
                
                # Upload do arquivo
                blob.upload_from_file(photo, content_type=photo.mimetype)
                
                # Tornar o blob público
                blob.make_public()
                
                # Obter a URL pública
                photo_url = blob.public_url
                user_data['photo_url'] = photo_url
            except Exception as e:
                logging.error(f'Erro ao fazer upload da foto: {str(e)}')
                flash('Erro ao fazer upload da foto. Por favor, tente novamente.', 'danger')
                return redirect(url_for('profile'))
        
        # Atualizar dados do usuário
        if display_name:
            user_data['display_name'] = display_name
        if phone:
            user_data['phone'] = phone
        
        # Atualizar dados no Firebase
        user_ref = db.reference(f'users/{user_id}')
        user_ref.update(user_data)
        flash('Perfil atualizado com sucesso!', 'success')
        return redirect(url_for('profile'))
    
    plan, plan_expiry, _, _ = get_user_plan_info(user_id)
    plan_data = PLANS.get(plan, PLANS['free'])
    
    # Combinar dados do usuário
    user_info_data = {
        'display_name': user_info.display_name or '',
        'email': user_info.email or '',
        'photo_url': user_info.photo_url or '',
        'phone': user_data.get('phone', '')
    }
    
    # Sobrescrever com dados do user_data quando disponíveis
    user_info_data.update({
        'display_name': user_data.get('display_name', user_info_data['display_name']),
        'photo_url': user_data.get('photo_url', user_info_data['photo_url']),
        'phone': user_data.get('phone', user_info_data['phone'])
    })
    
    return render_template('profile.html',
                         user_display_name=user_info_data['display_name'],
                         user_email=user_info_data['email'],
                         user_phone=user_info_data['phone'],
                         user_photo=user_info_data['photo_url'],
                         plan_data=plan_data,
                         plan_expiry=plan_expiry)

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    user_id = get_user_id()
    prefs = db.reference(f'user_prefs/{user_id}').get() or {}
    
    if request.method == 'POST':
        # Atualizar preferências
        new_prefs = request.form.to_dict()
        db.reference(f'user_prefs/{user_id}').update(new_prefs)
        flash('Preferências atualizadas com sucesso!', 'success')
        return redirect(url_for('settings'))
    
    # Carregar dados do usuário
    user_info = auth.get_user(user_id)
    user_data = db.reference(f'users/{user_id}').get() or {}
    
    # Carregar planos disponíveis e plano atual
    available_plans = PLANS  # Usando a constante PLANS diretamente
    current_plan_info = get_user_plan_info(user_id)
    current_plan = current_plan_info[0]  # Pega o nome do plano
    
    return render_template('settings.html',
                         user_info=user_info,
                         user_data=user_data,
                         prefs=prefs,
                         plans=available_plans,
                         current_plan=current_plan,
                         current_plan_info=current_plan_info)

@app.route('/brief/<brief_id>')
@login_required
def view_brief(brief_id):
    user_id = get_user_id()
    brief_ref = db.reference(f'briefs/{user_id}/{brief_id}')
    brief = brief_ref.get()
    if not brief:
        flash('Brief não encontrado.', 'danger')
        return redirect(url_for('my_briefs'))
    
    # Obter informações do plano do usuário
    plan, _, _, user_data = get_user_plan_info(user_id)
    
    brief['id'] = brief_id
    return render_template('view_brief.html', brief=brief, plan_data=PLANS.get(plan, PLANS['free']))

@app.route('/brief/<brief_id>/edit', methods=['GET', 'POST'])
@login_required
def edit_brief(brief_id):
    user_id = get_user_id()
    brief = db.reference(f'briefs/{user_id}/{brief_id}').get()
    if not brief:
        flash('Brief não encontrado.', 'danger')
        return redirect(url_for('my_briefs'))
    brief['id'] = brief_id
    
    # Obter informações do usuário
    user_ref = db.reference(f'users/{user_id}')
    user_data = user_ref.get()
    
    if request.method == 'POST':
        # Processar campos do formulário
        brief_data = {
            'brief_type': request.form.get('brief_type', brief['brief_type']),
            'template': request.form.get('template', brief['template']),
            'text': request.form.get('text', brief['text']),
            'result': request.form.get('result', brief['result']),
            'status': request.form.get('status', brief.get('status', 'concluido'))
        }
        
        # Processar campos extras
        extra_names = request.form.getlist('extra_name[]')
        extra_values = request.form.getlist('extra_value[]')
        extras = []
        for name, value in zip(extra_names, extra_values):
            if name and value:
                extras.append({'name': name, 'value': value})
        brief_data['extras'] = extras
        
        # Processar imagem
        if 'image' in request.files:
            image_file = request.files['image']
            if image_file.filename:
                # Salvar temporariamente
                temp_image_path = os.path.join(tempfile.gettempdir(), 'temp_image.jpg')
                image_file.save(temp_image_path)
                
                # Upload para o Firebase Storage
                bucket = storage.bucket()
                image_blob = bucket.blob(f'images/{uuid.uuid4()}.jpg')
                image_blob.upload_from_filename(temp_image_path)
                image_blob.make_public()
                brief_data['image_url'] = image_blob.public_url
                
                # Limpar arquivo temporário
                os.remove(temp_image_path)
        
        # Atualizar o brief no Firebase
        db.reference(f'briefs/{user_id}/{brief_id}').update(brief_data)
        
        flash('Brief atualizado com sucesso!', 'success')
        return redirect(url_for('view_brief', brief_id=brief_id))
    
    return render_template('edit_brief.html', user=user_data, brief=brief)

@app.route('/brief/<brief_id>/delete', methods=['POST'])
@login_required
def delete_brief(brief_id):
    user_id = get_user_id()
    try:
        db.reference(f'briefs/{user_id}/{brief_id}').delete()
        flash('Brief excluído com sucesso!', 'success')
    except Exception as e:
        flash(f'Erro ao excluir brief: {str(e)}', 'danger')
    return redirect(url_for('my_briefs'))

@app.context_processor
def inject_year():
    return {'year': datetime.now().year}

@app.context_processor
def inject_user_info():
    user_theme = 'light'
    user_photo = None
    user_plan = 'free'
    user_id = get_user_id()
    plan_data = PLANS['free']  # Define o plano padrão como 'free'
    
    if user_id:
        prefs = db.reference(f'user_prefs/{user_id}').get() or {}
        user_theme = prefs.get('theme', 'light')
        user_info = auth.get_user_by_email(session['user_email'])
        user_photo = prefs.get('photo_url', user_info.photo_url)
        user_data = db.reference(f'users/{user_id}').get() or {}
        user_plan = user_data.get('plan', 'free')
        plan_data = PLANS.get(user_plan, PLANS['free'])
    
    return dict(
        user_theme=user_theme, 
        user_photo=user_photo, 
        user_plan=user_plan, 
        plan_data=plan_data,
        plans=PLANS
    )

@app.route('/api/generate-brief', methods=['POST'])
def api_generate_brief():
    if not request.is_json:
        return jsonify({'error': 'Content-Type deve ser application/json'}), 400

    try:
        data = request.get_json()
        brief_id = data.get('brief_id')  # Novo campo para identificar brief existente
        text = data.get('text', '')
        brief_type = data.get('type')
        template = data.get('template')
        extras = data.get('extras', [])
        image_data = data.get('image')
        audio_data = data.get('audio')
        
        # Se brief_id foi fornecido, estamos editando um brief existente
        if brief_id:
            user_id = get_user_id()
            brief_ref = db.reference(f'briefs/{user_id}/{brief_id}')
            brief = brief_ref.get()
            if not brief:
                return jsonify({'error': 'Brief não encontrado'}), 404
            
            # Atualizar apenas os campos que foram fornecidos
            updates = {}
            if text:
                updates['text'] = text
            if brief_type:
                updates['brief_type'] = brief_type
            if template:
                updates['template'] = template
            if extras:
                updates['extras'] = extras
            
            # Se houver nova imagem ou áudio, processar e atualizar
            if image_data:
                # Processar imagem (código existente)
                # ...
                updates['image_url'] = image_url
                updates['image_analysis'] = image_analysis
            
            if audio_data:
                # Processar áudio (código existente)
                # ...
                updates['audio_url'] = audio_url
                updates['audio_text'] = audio_text
            
            # Atualizar o brief no Firebase
            brief_ref.update(updates)
            
            # Recuperar o brief atualizado
            brief = brief_ref.get()
            return jsonify({'result': brief})

        # Se não tem brief_id, estamos criando um novo brief (código existente)

        # Validar campos obrigatórios
        if not brief_type or not template:
            return jsonify({'error': 'Tipo e template são obrigatórios'}), 400

        # Processar áudio se fornecido
        audio_text = ''
        audio_url = ''
        if audio_data:
            try:
                # Decodificar o áudio base64
                audio_content = base64.b64decode(audio_data.split(',')[1])
                
                # Salvar temporariamente
                temp_audio_path = os.path.join(tempfile.gettempdir(), 'temp_audio.m4a')
                with open(temp_audio_path, 'wb') as f:
                    f.write(audio_content)

                # Transcrever usando a API do Groq
                with open(temp_audio_path, 'rb') as f:
                    transcription = groq_client.audio.transcriptions.create(
                        file=(temp_audio_path, f.read()),
                        model="whisper-large-v3-turbo",
                        response_format="verbose_json"
                    )
                    audio_text = transcription.text

                # Upload para o Firebase Storage
                bucket = storage.bucket()
                audio_blob = bucket.blob(f'audio/{uuid.uuid4()}.m4a')
                audio_blob.upload_from_filename(temp_audio_path)
                audio_blob.make_public()
                audio_url = audio_blob.public_url

                # Limpar arquivo temporário
                os.remove(temp_audio_path)
            except Exception as e:
                print(f"Erro ao processar áudio: {str(e)}")
                return jsonify({'error': 'Erro ao processar áudio'}), 500

        # Processar imagem se fornecida
        image_url = ''
        image_analysis = ''
        if image_data:
            try:
                # Decodificar a imagem base64
                image_content = base64.b64decode(image_data.split(',')[1])
                
                # Salvar temporariamente
                temp_image_path = os.path.join(tempfile.gettempdir(), 'temp_image.jpg')
                with open(temp_image_path, 'wb') as f:
                    f.write(image_content)

                # Upload para o Firebase Storage
                bucket = storage.bucket()
                image_blob = bucket.blob(f'images/{uuid.uuid4()}.jpg')
                image_blob.upload_from_filename(temp_image_path)
                image_blob.make_public()
                image_url = image_blob.public_url

                # Analisar imagem usando o modelo maverick
                with open(temp_image_path, 'rb') as f:
                    image_analysis = groq_client.chat.completions.create(
                        model="meta-llama/llama-4-maverick-17b-128e-instruct",
                        messages=[
                            {
                                "role": "system",
                                "content": "Você é um assistente especializado em análise de imagens para briefings. Analise a imagem fornecida e extraia informações relevantes para um briefing profissional."
                            },
                            {
                                "role": "user",
                                "content": [
                                    {
                                        "type": "text",
                                        "text": "Analise esta imagem e extraia informações relevantes para um briefing profissional. Inclua detalhes sobre cores, composição, elementos visuais e qualquer outro aspecto importante."
                                    },
                                    {
                                        "type": "image_url",
                                        "image_url": {
                                            "url": image_url
                                        }
                                    }
                                ]
                            }
                        ]
                    ).choices[0].message.content

                # Limpar arquivo temporário
                os.remove(temp_image_path)
            except Exception as e:
                print(f"Erro ao processar imagem: {str(e)}")
                return jsonify({'error': 'Erro ao processar imagem'}), 500

        # Construir o prompt combinando todas as informações
        prompt = f"""Crie um briefing profissional baseado nas seguintes informações:

Tipo de Briefing: {brief_type}
Template: {template}

Texto do Cliente:
{text}

Transcrição do Áudio:
{audio_text}

Análise da Imagem:
{image_analysis}

Campos Extras:
{json.dumps(extras, indent=2, ensure_ascii=False)}

Por favor, crie um briefing detalhado e profissional que incorpore todas essas informações de forma coesa e estruturada."""

        # Gerar o briefing usando o modelo scout
        completion = groq_client.chat.completions.create(
            model="meta-llama/llama-4-scout-17b-16e-instruct",
            messages=[
                {
                    "role": "system",
                    "content": "Você é um assistente especializado em criar briefings profissionais. Seu objetivo é transformar as informações fornecidas em um briefing bem estruturado e detalhado."
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
        )

        result = completion.choices[0].message.content

        # Salvar no Firebase
        user_id = get_user_id()
        if user_id:
            brief_data = {
                'text': text,
                'audio_text': audio_text,
                'audio_url': audio_url,
                'image_url': image_url,
                'image_analysis': image_analysis,
                'brief_type': brief_type,
                'template': template,
                'extras': [{'name': extra['name'], 'value': extra['value']} for extra in extras],
                'result': result,
                'created_at': datetime.utcnow().isoformat() + 'Z',
                'status': 'concluido',
            }
            brief_ref = db.reference(f'briefs/{user_id}').push(brief_data)
            brief_data['id'] = brief_ref.key

        return jsonify({'result': result, 'brief_id': brief_ref.key})
    except Exception as e:
        print(f"Erro ao gerar briefing: {str(e)}")
        return jsonify({'error': f'Erro ao gerar briefing: {str(e)}'}), 500
def export_brief_pdf(brief_id):
    if 'user' not in session or 'user_email' not in session:
        flash('Faça login para acessar seus briefs.', 'warning')
        return redirect(url_for('login'))
    user_id = get_user_id()
    brief = db.reference(f'briefs/{user_id}/{brief_id}').get()
    if not brief:
        flash('Brief não encontrado.', 'danger')
        return redirect(url_for('my_briefs'))

    # Geração do PDF em memória
    from io import BytesIO
    buffer = BytesIO()
    p = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    y = height - 50
    p.setFont('Helvetica-Bold', 16)
    p.drawString(50, y, f"Brief: {brief.get('title', 'Sem título')}")
    y -= 30
    p.setFont('Helvetica', 12)
    for key, value in brief.items():
        if key == 'id':
            continue
        if isinstance(value, str):
            lines = value.split('\n')
            for line in lines:
                if y < 50:
                    p.showPage()
                    y = height - 50
                    p.setFont('Helvetica', 12)
                p.drawString(50, y, f"{key.capitalize()}: {line}" if line else "")
                y -= 18
        else:
            if y < 50:
                p.showPage()
                y = height - 50
                p.setFont('Helvetica', 12)

@app.route('/brief/<brief_id>/export/docx')
@login_required
def export_brief_docx(brief_id):
    user_id = get_user_id()
    brief = db.reference(f'briefs/{user_id}/{brief_id}').get()
    if not brief:
        flash('Brief não encontrado.', 'danger')
        return redirect(url_for('my_briefs'))

    # Criar documento DOCX
    doc = Document()
    
    # Estilo do título
    title_style = doc.styles['Heading 1']
    title_style.font.size = Pt(32)
    title_style.font.bold = True
    title_style.font.color.rgb = RGBColor(44, 62, 80)  # Cor azul escuro
    
    # Estilo dos cabeçalhos
    heading_style = doc.styles['Heading 2']
    heading_style.font.size = Pt(18)
    heading_style.font.bold = True
    heading_style.font.color.rgb = RGBColor(52, 152, 219)  # Cor azul
    
    # Estilo dos campos
    field_style = doc.styles.add_style('Field', WD_STYLE_TYPE.PARAGRAPH)
    field_style.font.size = Pt(14)
    field_style.font.color.rgb = RGBColor(84, 153, 199)  # Cor azul claro
    
    # Adicionar título
    title = doc.add_heading('Briefing', 0)
    title.alignment = 1  # Centralizado
    
    # Adicionar informações básicas
    doc.add_heading('Informações do Briefing', level=1)
    doc.add_paragraph(f'Tipo: {brief.get("brief_type", "Não especificado")}', style='Field')
    doc.add_paragraph(f'Template: {brief.get("template", "Não especificado")}', style='Field')
    doc.add_paragraph(f'Data: {brief.get("created_at", "Não especificada")}', style='Field')
    
    # Adicionar descrição
    doc.add_heading('Descrição do Projeto', level=1)
    doc.add_paragraph(brief.get('text', 'Nenhuma descrição fornecida'))
    
    # Adicionar resultado
    doc.add_heading('Resultado do Briefing', level=1)
    doc.add_paragraph(brief.get('result', 'Nenhum resultado gerado'))
    
    # Adicionar campos extras
    if brief.get('extras'):
        doc.add_heading('Campos Adicionais', level=1)
        for extra in brief['extras']:
            doc.add_paragraph(f"{extra.get('name', '')}: {extra.get('value', '')}", style='Field')
    
    # Adicionar status
    doc.add_heading('Status', level=1)
    status_paragraph = doc.add_paragraph()
    status_run = status_paragraph.add_run(brief.get('status', 'Não especificado'))
    status_run.bold = True
    status_run.font.color.rgb = RGBColor(39, 174, 96)  # Cor verde
    
    # Adicionar imagem se existir
    if brief.get('image_url'):
        try:
            import requests
            from docx.shared import Inches
            response = requests.get(brief['image_url'])
            if response.status_code == 200:
                from io import BytesIO
                image_data = BytesIO(response.content)
                doc.add_heading('Imagem de Referência', level=1)
                doc.add_picture(image_data, width=Inches(6))
        except Exception as e:
            print(f"Erro ao adicionar imagem: {str(e)}")
    
    # Salvar documento temporariamente
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.docx')
    doc.save(temp_file.name)
    
    # Enviar arquivo
    return send_file(
        temp_file.name,
        as_attachment=True,
        download_name=f'brief_{brief_id}.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/brief/<brief_id>/export/pdf')
@login_required
def export_brief_pdf(brief_id):
    user_id = get_user_id()
    brief = db.reference(f'briefs/{user_id}/{brief_id}').get()
    if not brief:
        flash('Brief não encontrado.', 'danger')
        return redirect(url_for('my_briefs'))
    
    # Cores para formatação
    TEXT_COLOR = (0.4, 0.4, 0.4)  # Cinza escuro
    SECONDARY_COLOR = (0.2, 0.4, 0.5)  # Azul médio
    SUCCESS_COLOR = (0.16, 0.8, 0.44)  # Verde
    
    # Configurações de página
    PAGE_WIDTH = 595  # 8.27 polegadas
    PAGE_HEIGHT = 842  # 11.69 polegadas
    margin = 50
    
    # Criar arquivo temporário
    temp_file = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
    c = canvas.Canvas(temp_file.name, pagesize=(PAGE_WIDTH, PAGE_HEIGHT))
    
    # Função para adicionar texto com quebra de linha e formatação
    def add_text(text, y, font_size=12, line_height=16, max_width=500, color=TEXT_COLOR, bold=False):
        c.setFont('Helvetica-Bold' if bold else 'Helvetica', font_size)
        c.setFillColorRGB(*color)
        if not text.strip():
            return y
            
        # Separar texto em parágrafos
        paragraphs = text.split('\n\n')
        
        for paragraph in paragraphs:
            # Separar linhas
            lines = paragraph.split('\n')
            for line in lines:
                # Separar em palavras
                words = line.split()
                if not words:
                    y -= line_height  # Espaço para linha em branco
                    continue
                    
                current_line = words[0]
                for word in words[1:]:
                    if c.stringWidth(current_line + ' ' + word) < max_width:
                        current_line += ' ' + word
                    else:
                        if y < margin + 20:  # Se estiver chegando perto do fim da página
                            c.showPage()  # Quebra de página
                            y = PAGE_HEIGHT - margin  # Reinicia na nova página
                            
                        c.drawString(margin, y, current_line)
                        y -= line_height
                        current_line = word
                
                if y < margin + 20:  # Se estiver chegando perto do fim da página
                    c.showPage()  # Quebra de página
                    y = PAGE_HEIGHT - margin  # Reinicia na nova página
                    
                c.drawString(margin, y, current_line)
                y -= line_height
            
            # Espaço adicional entre parágrafos
            if y < margin + 20:  # Se estiver chegando perto do fim da página
                c.showPage()  # Quebra de página
                y = PAGE_HEIGHT - margin  # Reinicia na nova página
                
            y -= line_height
        
        return y
    
    # Função para adicionar seção com título
    def add_section(title, content, y, font_size=12, line_height=16):
        # Verificar se é um título principal (mais importante)
        main_titles = [
            'Briefing de Marketing:', 'Introdução:', 'Análise da Imagem:',
            'Objetivos da Loja:', 'Público-Alvo:', 'Requisitos da Loja:',
            'Produtos a Serem Oferecidos:', 'Estratégia de Marketing:',
            'Campos Extras:', 'Conclusão:'
        ]
        
        # Verificar se estamos chegando perto do fim da página
        if y < margin + 60:  # Se estiver muito perto do fim
            c.showPage()  # Quebra de página
            y = PAGE_HEIGHT - margin  # Reinicia na nova página
        
        # Configurar estilo do título
        if any(title.startswith(t) for t in main_titles):
            c.setFont('Helvetica-Bold', 16)
            c.setFillColorRGB(*SECONDARY_COLOR)
            y -= 40  # Espaço maior para títulos principais
        else:
            c.setFont('Helvetica-Bold', 14)
            c.setFillColorRGB(*SECONDARY_COLOR)
            y -= 30
        
        # Desenhar título
        c.drawString(margin, y, title)
        y -= 20
        
        # Adicionar conteúdo se houver
        if content:
            # Verificar se o conteúdo tem listas (começa com * ou +)
            if content.strip().startswith('* ') or content.strip().startswith('+ '):
                # Para listas, usar fonte menor e espaçamento mais apertado
                y = add_text(content, y, font_size=11, line_height=14)
            else:
                y = add_text(content, y, font_size, line_height)
            y -= 20  # Espaço após o conteúdo
        
        return y
    
    # Função para adicionar imagem
    def add_image(image_path, y, max_width=500, max_height=300):
        try:
            # Obter dimensões originais
            image = Image.open(image_path)
            width, height = image.size
            
            # Calcular proporção
            ratio = min(max_width/width, max_height/height)
            
            # Calcular novas dimensões mantendo proporção
            new_width = int(width * ratio)
            new_height = int(height * ratio)
            
            # Desenhar imagem
            c.drawImage(image_path, margin, y - new_height, width=new_width, height=new_height, preserveAspectRatio=True)
            y -= new_height + 30  # Espaço para a imagem e abaixo dela
            return y
        except Exception as e:
            print(f"Erro ao processar imagem: {str(e)}")
            return y
    
    # Posição inicial
    y = PAGE_HEIGHT - margin
    
    # Informações do Briefing
    y = add_section('Informações do Briefing', '', y)
    y = add_text(f"Tipo: {brief.get('brief_type', 'Não especificado')}", y, font_size=12, bold=True)
    y = add_text(f"Template: {brief.get('template', 'Não especificado')}", y, font_size=12, bold=True)
    y = add_text(f"Data: {brief.get('created_at', 'Não especificada')}", y, font_size=12, bold=True)
    y -= 30  # Espaço adicional após informações
    
    # Descrição do Projeto
    y = add_section('Descrição do Projeto', brief.get('text', 'Nenhuma descrição fornecida'), y)
    
    # Resultado do Briefing
    y = add_section('Resultado do Briefing', brief.get('result', 'Nenhum resultado gerado'), y)
    
    # Imagem de Referência
    if brief.get('image_url'):
        try:
            image = requests.get(brief.get('image_url'))
            if image.status_code == 200:
                image_path = tempfile.NamedTemporaryFile(delete=False, suffix='.jpg').name
                with open(image_path, 'wb') as f:
                    f.write(image.content)
                y = add_image(image_path, y)
        except Exception as e:
            print(f"Erro ao processar imagem: {str(e)}")
    
    # Campos Adicionais
    if brief.get('extras'):
        y = add_section('Campos Adicionais', '', y)
        for extra in brief.get('extras', []):
            y = add_text(f"{extra.get('name', '')}: {extra.get('value', '')}", y, font_size=12, bold=True)
        y -= 20  # Espaço adicional após campos adicionais
    
    # Status
    y = add_section('Status', '', y)
    status = brief.get('status', 'Não especificado')
    c.setFont('Helvetica-Bold', 16)
    c.setFillColorRGB(*SUCCESS_COLOR)
    c.drawString(margin, y, status)
    
    # Finalizar e salvar PDF
    c.save()
    
    # Enviar arquivo
    return send_file(
        temp_file.name,
        as_attachment=True,
        download_name=f'brief_{brief_id}.pdf',
        mimetype='application/pdf'
    )

@app.route('/brief/<brief_id>/export/txt')
def export_brief_txt(brief_id):
    user_id = get_user_id()
    brief = db.reference(f'briefs/{user_id}/{brief_id}').get()
    if not brief:
        flash('Brief não encontrado.', 'danger')
        return redirect(url_for('my_briefs'))

    # Criar conteúdo TXT formatado
    lines = []
    
    # Título centralizado
    title = "Briefing"
    lines.append("".join("=" * 80))
    lines.append(title.center(80))
    lines.append("".join("=" * 80))
    lines.append("")
    
    # Informações do Briefing
    lines.append("INFORMAÇÕES DO BRIEFING")
    lines.append("".join("-" * 80))
    lines.append(f"Tipo: {brief.get('brief_type', 'Não especificado')}")
    lines.append(f"Template: {brief.get('template', 'Não especificado')}")
    lines.append(f"Data: {brief.get('created_at', 'Não especificada')}")
    lines.append("")
    
    # Descrição do Projeto
    lines.append("DESCRIÇÃO DO PROJETO")
    lines.append("".join("-" * 80))
    text = brief.get('text', 'Nenhuma descrição fornecida')
    # Quebra de linha para melhor legibilidade
    text_lines = text.split('\n')
    for line in text_lines:
        lines.append(f"  {line}")
    lines.append("")
    
    # Resultado do Briefing
    lines.append("RESULTADO DO BRIEFING")
    lines.append("".join("-" * 80))
    result = brief.get('result', 'Nenhum resultado gerado')
    # Quebra de linha para melhor legibilidade
    result_lines = result.split('\n')
    for line in result_lines:
        lines.append(f"  {line}")
    lines.append("")
    
    # Campos Extras
    if brief.get('extras'):
        lines.append("CAMPOS ADICIONAIS")
        lines.append("".join("-" * 80))
        for extra in brief.get('extras', []):
            lines.append(f"  {extra.get('name', '')}:")
            lines.append(f"    {extra.get('value', '')}")
        lines.append("")
    
    # Status
    lines.append("STATUS")
    lines.append("".join("-" * 80))
    lines.append(f"  {brief.get('status', 'Não especificado')}")
    
    # Juntar todas as linhas
    txt_content = '\n'.join(lines)
    
    from io import BytesIO
    buffer = BytesIO(txt_content.encode('utf-8'))
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"brief_{brief_id}.txt",
        mimetype='text/plain'
    )

@app.route('/brief/<brief_id>/export/html')
@login_required
def export_brief_html(brief_id):
    user_id = get_user_id()
    brief = db.reference(f'briefs/{user_id}/{brief_id}').get()
    if not brief:
        flash('Brief não encontrado.', 'danger')
        return redirect(url_for('my_briefs'))
    
    # Criar HTML com formatação
    html_content = """
    <!DOCTYPE html>
    <html lang="pt-BR">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Briefing</title>
        <style>
            body {
                font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
                max-width: 800px;
                margin: 0 auto;
                padding: 20px;
                line-height: 1.8;
                color: #333;
                background-color: #f5f5f5;
            }
            .important {
                font-weight: bold;
                color: #2c3e50;
            }
            .container {
                background-color: white;
                padding: 30px;
                border-radius: 10px;
                box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
            }
            h1 {
                color: #2c3e50;
                text-align: center;
                margin-bottom: 30px;
                font-size: 2.5em;
            }
            h2 {
                color: #3498db;
                margin-top: 20px;
                margin-bottom: 15px;
                font-size: 1.5em;
                border-bottom: 2px solid #3498db;
                padding-bottom: 5px;
            }
            .section {
                margin-bottom: 30px;
            }
            .field {
                margin-bottom: 15px;
                display: flex;
                align-items: flex-start;
                gap: 10px;
            }
            .field-name {
                font-weight: bold;
                color: #2c3e50;
                min-width: 120px;
                font-size: 1.1em;
            }
            .field-value {
                flex: 1;
                color: #555;
            }
            .status {
                font-weight: bold;
                color: #27ae60;
                font-size: 1.2em;
            }
            .date {
                color: #7f8c8d;
                font-size: 0.9em;
            }
            .image-container {
                margin: 20px 0;
                text-align: center;
            }
            .image-container img {
                max-width: 100%;
                height: auto;
                border-radius: 8px;
                box-shadow: 0 2px 10px rgba(0, 0, 0, 0.1);
            }
            .result {
                white-space: pre-wrap;
                line-height: 1.6;
            }
        </style>
    </head>
    <body>
        <div class="container">
            <h1>Briefing</h1>
            <div class="section">
                <h2 class="important">Informações do Briefing</h2>
                <div class="field">
                    <span class="field-name important">Tipo:</span>
                    <span class="field-value">{{ brief.get('brief_type', 'Não especificado') }}</span>
                </div>
                <div class="field">
                    <span class="field-name important">Template:</span>
                    <span class="field-value">{{ brief.get('template', 'Não especificado') }}</span>
                </div>
                <div class="field">
                    <span class="field-name important">Data:</span>
                    <span class="field-value date">{{ brief.get('created_at', 'Não especificada') }}</span>
                </div>
            </div>
            
            <div class="section">
                <h2 class="important">Descrição do Projeto</h2>
                <p class="field-value">{{ brief.get('text', 'Nenhuma descrição fornecida') }}</p>
            </div>
            
            <div class="section">
                <h2 class="important">Resultado do Briefing</h2>
                <div class="result">{{ brief.get('result', 'Nenhum resultado gerado') }}</div>
            </div>
            
            {% if brief.get('image_url') %}
            <div class="section">
                <h2 class="important">Imagem de Referência</h2>
                <div class="image-container">
                    <img src="{{ brief.get('image_url') }}" alt="Imagem de referência">
                </div>
            </div>
            {% endif %}
            
            {% if brief.get('extras') %}
            <div class="section">
                <h2 class="important">Campos Adicionais</h2>
                {% for extra in brief.get('extras', []) %}
                <div class="field">
                    <span class="field-name important">{{ extra.get('name', '') }}:</span>
                    <span class="field-value">{{ extra.get('value', '') }}</span>
                </div>
                {% endfor %}
            </div>
            {% endif %}
            
            <div class="section">
                <h2 class="important">Status</h2>
                <p class="status">{{ brief.get('status', 'Não especificado') }}</p>
            </div>
        </div>
    </body>
    </html>
    """
    
    return Response(
        render_template_string(html_content, brief=brief),
        mimetype='text/html',
        headers={
            'Content-Disposition': f'attachment; filename="brief_{brief_id}.html"'
        }
    )

@app.route('/plans')
def plans():
    if 'user' not in session or 'user_email' not in session:
        flash('Faça login para acessar os planos.', 'warning')
        return redirect(url_for('login'))
    user_id = get_user_id()
    plan, plan_expiry, briefs_this_month, user_data = get_user_plan_info(user_id)
    return render_template('plans.html', plans=PLANS, user_plan=plan, plan_expiry=plan_expiry, briefs_this_month=briefs_this_month)

@app.route('/upgrade/<plan_key>', methods=['POST'])
def upgrade(plan_key):
    logging.info(f'[Upgrade] Iniciando processo de upgrade para o plano {plan_key}')
    
    if plan_key not in PLANS:
        logging.error(f'[Upgrade] Plano inválido: {plan_key}')
        return jsonify({'error': 'Plano inválido.'}), 400
        
    if 'user' not in session or 'user_email' not in session:
        logging.error('[Upgrade] Usuário não está logado')
        return jsonify({'error': 'Login necessário.'}), 401
    
    user_id = get_user_id()
    logging.info(f'[Upgrade] User ID: {user_id}')
    
    # Se for o plano gratuito, atualiza direto no banco
    if plan_key == 'free':
        try:
            db.reference(f'users/{user_id}').update({
                'plan': 'free',
                'plan_expiry': None
            })
            flash('Plano atualizado com sucesso!', 'success')
            return redirect(url_for('dashboard'))
        except Exception as e:
            logging.error(f'[Upgrade] Erro ao atualizar para plano gratuito: {str(e)}')
            flash('Erro ao atualizar plano. Por favor, tente novamente mais tarde.', 'error')
            return redirect(url_for('plans'))
    
    # Para planos pagos, usa o Stripe
    if not stripe.api_key:
        logging.error('[Upgrade] STRIPE_API_KEY não configurada')
        flash('Erro de configuração do sistema de pagamento. Por favor, tente novamente mais tarde.', 'error')
        return redirect(url_for('plans'))
    
    session['upgrade_plan'] = plan_key
    
    try:
        # Stripe Checkout Session
        domain_url = request.url_root.strip('/')
        price_id = PLANS[plan_key]['stripe_price_id']
        
        logging.info(f'[Upgrade] Domain URL: {domain_url}')
        logging.info(f'[Upgrade] Price ID: {price_id}')
        
        if not price_id:
            logging.error(f'[Upgrade] Stripe Price ID não configurado para o plano {plan_key}')
            flash('Erro de configuração do sistema de pagamento. Por favor, tente novamente mais tarde.', 'error')
            return redirect(url_for('plans'))
        
        # Verificar se o preço existe no Stripe
        try:
            price = stripe.Price.retrieve(price_id)
            logging.info(f'[Upgrade] Preço encontrado no Stripe: {price}')
        except Exception as e:
            logging.error(f'[Upgrade] Erro ao verificar preço no Stripe: {str(e)}')
            flash('Erro de configuração do sistema de pagamento. Por favor, tente novamente mais tarde.', 'error')
            return redirect(url_for('plans'))
        
        stripe_session = stripe.checkout.Session.create(
            payment_method_types=['card'],
            line_items=[{
                'price': price_id,
                'quantity': 1
            }],
            mode='subscription',  # Changed from 'payment' to 'subscription'
            success_url=f'{domain_url}/payment/success?session_id={{CHECKOUT_SESSION_ID}}',
            cancel_url=f'{domain_url}/plans',
            metadata={'user_id': user_id, 'plan': plan_key}
        )
        
        logging.info(f'[Upgrade] Sessão do Stripe criada com sucesso: {stripe_session.id}')
        return redirect(stripe_session.url)
    except stripe.error.StripeError as e:
        logging.error(f'[Upgrade] Erro do Stripe: {str(e)}')
        flash('Erro ao processar pagamento. Por favor, tente novamente mais tarde.', 'error')
        return redirect(url_for('plans'))
    except Exception as e:
        logging.error(f'[Upgrade] Erro inesperado: {str(e)}')
        flash('Erro ao processar pagamento. Por favor, tente novamente mais tarde.', 'error')
        return redirect(url_for('plans'))

@app.route('/payment/success')
def payment_success():
    user_id = get_user_id()
    session_id = request.args.get('session_id')
    if user_id and session_id:
        try:
            stripe_session = stripe.checkout.Session.retrieve(session_id)
            user_data = db.reference(f'users/{user_id}').get() or {}
            if user_data.get('plan') == stripe_session['metadata']['plan']:
                flash(f'Plano {PLANS[user_data["plan"]]["name"]} ativado com sucesso!', 'success')
            else:
                flash('Plano ainda não atualizado. Por favor, aguarde alguns instantes.', 'warning')
        except Exception as e:
            logging.error(f'[Payment Success] Erro ao verificar sessão: {e}')
            flash('Erro ao verificar pagamento. Entre em contato com o suporte.', 'error')
    return render_template('payment_success.html')

@app.route('/webhook/stripe', methods=['POST'])
def stripe_webhook():
    payload = request.data
    sig_header = request.headers.get('Stripe-Signature')
    endpoint_secret = os.getenv('STRIPE_WEBHOOK_SECRET')
    
    logging.info(f'[Stripe Webhook] Recebido evento. Headers: {dict(request.headers)}')
    logging.info(f'[Stripe Webhook] Payload: {payload.decode()}')
    logging.info(f'[Stripe Webhook] Signature: {sig_header}')
    logging.info(f'[Stripe Webhook] Secret configurado: {"Sim" if endpoint_secret else "Não"}')
    
    if not endpoint_secret:
        logging.error('[Stripe Webhook] STRIPE_WEBHOOK_SECRET não configurado')
        return 'Webhook secret not configured', 500
        
    try:
        event = stripe.Webhook.construct_event(payload, sig_header, endpoint_secret)
        logging.info(f'[Stripe Webhook] Evento construído com sucesso: {event["type"]}')
        logging.info(f'[Stripe Webhook] Dados do evento: {event["data"]}')
    except ValueError as e:
        logging.warning(f'[Stripe Webhook] Payload inválido: {e}')
        return 'Invalid payload', 400
    except stripe.error.SignatureVerificationError as e:
        logging.error(f'[Stripe Webhook] Falha na assinatura: {e}')
        return 'Invalid signature', 400
    except Exception as e:
        logging.error(f'[Stripe Webhook] Erro inesperado: {e}')
        return 'Webhook error', 400

    if event['type'] == 'checkout.session.completed':
        session_obj = event['data']['object']
        user_id = session_obj['metadata'].get('user_id')
        plan_key = session_obj['metadata'].get('plan')
        
        logging.info(f'[Stripe Webhook] Metadata: user_id={user_id}, plan_key={plan_key}')
        
        if not user_id or not plan_key:
            logging.error(f'[Stripe Webhook] user_id ou plan_key ausentes: user_id={user_id}, plan_key={plan_key}')
            return 'Missing user_id or plan_key', 400
            
        if plan_key not in PLANS:
            logging.error(f'[Stripe Webhook] Plano inválido: {plan_key}')
            return 'Invalid plan', 400
            
        try:
            expiry = (datetime.utcnow().replace(day=1) + timedelta(days=32)).replace(day=1)
            expiry_str = expiry.strftime('%Y-%m-%d')
            
            update_data = {
                'plan': plan_key,
                'plan_expiry': expiry_str,
                'stripe_subscription_id': session_obj.get('subscription'),
                'stripe_customer_id': session_obj.get('customer')
            }
            logging.info(f'[Stripe Webhook] Dados a serem atualizados: {update_data}')
            
            db.reference(f'users/{user_id}').update(update_data)
            
            updated_user_data = db.reference(f'users/{user_id}').get() or {}
            logging.info(f'[Stripe Webhook] Dados do usuário após atualização: {updated_user_data}')
            
            create_notification(
                user_id=user_id,
                title="Plano Atualizado",
                message=f"Seu plano foi atualizado para {PLANS[plan_key]['name']} com sucesso! Expira em {expiry_str}.",
                type="success"
            )
            
            email = session_obj.get('customer_email')
            if email:
                try:
                    send_email(
                        subject='Pagamento confirmado - EazyBrief',
                        recipients=[email],
                        body=f'Seu pagamento foi confirmado e seu plano {PLANS[plan_key]["name"]} já está ativo!'
                    )
                    logging.info(f'[Stripe Webhook] Email de confirmação enviado para {email}')
                except Exception as e:
                    logging.error(f'[Stripe Webhook] Erro ao enviar email: {e}')
        except Exception as e:
            logging.error(f'[Stripe Webhook] Erro ao atualizar plano: {e}')
            return 'Error updating plan', 500
    
    elif event['type'] == 'invoice.paid':
        invoice = event['data']['object']
        subscription_id = invoice['parent']['subscription_details']['subscription']
        customer_id = invoice['customer']
        
        try:
            subscription = stripe.Subscription.retrieve(subscription_id)
            checkout_session = stripe.checkout.Session.list(subscription=subscription_id, limit=1).data[0]
            user_id = checkout_session['metadata'].get('user_id')
            plan_key = checkout_session['metadata'].get('plan')
            
            if not user_id or not plan_key:
                logging.error(f'[Stripe Webhook] user_id ou plan_key ausentes em invoice.paid: user_id={user_id}, plan_key={plan_key}')
                return 'Missing user_id or plan_key', 400
                
            if plan_key not in PLANS:
                logging.error(f'[Stripe Webhook] Plano inválido em invoice.paid: {plan_key}')
                return 'Invalid plan', 400
                
            expiry = (datetime.utcnow().replace(day=1) + timedelta(days=32)).replace(day=1)
            expiry_str = expiry.strftime('%Y-%m-%d')
            
            update_data = {
                'plan': plan_key,
                'plan_expiry': expiry_str,
                'stripe_subscription_id': subscription_id,
                'stripe_customer_id': customer_id
            }
            logging.info(f'[Stripe Webhook] Atualizando plano para invoice.paid: {update_data}')
            
            db.reference(f'users/{user_id}').update(update_data)
            
            create_notification(
                user_id=user_id,
                title="Plano Atualizado",
                message=f"Seu plano foi atualizado para {PLANS[plan_key]['name']} com sucesso! Expira em {expiry_str}.",
                type="success"
            )
            
            email = invoice.get('customer_email')
            if email:
                send_email(
                    subject='Pagamento confirmado - EazyBrief',
                    recipients=[email],
                    body=f'Seu pagamento foi confirmado e seu plano {PLANS[plan_key]["name"]} já está ativo!'
                )
        except Exception as e:
            logging.error(f'[Stripe Webhook] Erro ao processar invoice.paid: {e}')
            return 'Error processing invoice.paid', 500
    
    else:
        logging.info(f'[Stripe Webhook] Evento ignorado: {event["type"]}')
    
    return '', 200
# Lista de e-mails admin
ADMINS = [
    'robbiealgon@gmail.com',
    'robertocahimagonga@gmail.com'
]

def is_admin(email):
    return email and email.lower() in ADMINS
def admin_dashboard():
    if 'user_email' not in session or not is_admin(session['user_email']):
        flash('Acesso restrito ao administrador.', 'danger')
        return redirect(url_for('index'))
    # Buscar usuários e métricas
    users_ref = db.reference('users')
    users_data = users_ref.get() or {}
    total_users = len(users_data)
    active_pro = sum(1 for u in users_data.values() if u.get('plan') == 'pro')
    active_premium = sum(1 for u in users_data.values() if u.get('plan') == 'premium')
    briefs_this_month = sum(
        list(u.get('briefs_used', {}).values())[-1] if u.get('briefs_used') else 0
        for u in users_data.values()
    )
    
    # Obter informações do plano do usuário
    plan, _, _, _ = get_user_plan_info(session['user_id'])
    plan_data = PLANS.get(plan, PLANS['free'])
    
    return render_template(
        'admin.html',
        total_users=total_users,
        active_pro=active_pro,
        active_premium=active_premium,
        briefs_this_month=briefs_this_month,
        users=users_data,
        PLANS=PLANS,
        plan_data=plan_data
    )
def inject_ga_measurement_id():
    return {'GA_MEASUREMENT_ID': os.getenv('GA_MEASUREMENT_ID', '')}

@app.context_processor
def inject_stripe_key():
    return {'stripe_publishable_key': stripe_publishable_key}

# Função para criar notificação
def create_notification(user_id, title, message, type='info'):
    notification = {
        'id': str(uuid.uuid4()),
        'title': title,
        'message': message,
        'type': type,
        'read': False,
        'created_at': datetime.utcnow().isoformat() + 'Z'
    }
    db.reference(f'notifications/{user_id}').push(notification)
    return notification

@app.route('/notifications')
@login_required
def notifications():
    user_id = get_user_id()
    notifications_ref = db.reference(f'notifications/{user_id}')
    notifications_data = notifications_ref.get() or {}
    
    notifications = []
    for key, val in notifications_data.items():
        val['id'] = key
        notifications.append(val)
    
    notifications.sort(key=lambda x: x['created_at'], reverse=True)
    return render_template('notifications.html', notifications=notifications)

@app.route('/api/notifications/<notification_id>/read', methods=['POST'])
@login_required
def mark_notification_read(notification_id):
    user_id = get_user_id()
    try:
        db.reference(f'notifications/{user_id}/{notification_id}').update({'read': True})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

@app.route('/api/notifications/read-all', methods=['POST'])
@login_required
def mark_all_notifications_read():
    user_id = get_user_id()
    try:
        notifications_ref = db.reference(f'notifications/{user_id}')
        notifications = notifications_ref.get() or {}
        for key in notifications:
            notifications_ref.child(key).update({'read': True})
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)})

# Adicionar notificação ao menu lateral
@app.context_processor
def inject_notification_count():
    notification_count = 0
    if 'user' in session and 'user_email' in session:
        user_id = get_user_id()
        if user_id:
            notifications_ref = db.reference(f'notifications/{user_id}')
            notifications = notifications_ref.get() or {}
            notification_count = sum(1 for n in notifications.values() if not n.get('read', False))
    return dict(notification_count=notification_count)

@app.route('/support')
@login_required
def support():
    user_id = get_user_id()
    plan, _, _, _ = get_user_plan_info(user_id)
    plan_data = PLANS.get(plan, PLANS['free'])
    return render_template('support.html', plan_data=plan_data)

@app.route('/api/contact', methods=['POST'])
def contact():
    if not request.is_json:
        return jsonify({'success': False, 'error': 'Content-Type deve ser application/json'}), 400
    
    data = request.get_json()
    subject = data.get('subject')
    message = data.get('message')
    
    if not subject or not message:
        return jsonify({'success': False, 'error': 'Assunto e mensagem são obrigatórios'}), 400
    
    try:
        # Salvar mensagem no Firebase
        contact_data = {
            'subject': subject,
            'message': message,
            'email': session.get('user_email', 'Não logado'),
            'created_at': datetime.utcnow().isoformat() + 'Z',
            'status': 'new'
        }
        db.reference('contact_messages').push(contact_data)
        
        # Enviar email para os admins
        admin_emails = ADMINS
        for admin_email in admin_emails:
            try:
                send_email(
                    subject=f'Nova mensagem de contato: {subject}',
                    recipients=[admin_email],
                    body=f"""
                    Nova mensagem de contato recebida:
                    
                    De: {session.get('user_email', 'Não logado')}
                    Assunto: {subject}
                    
                    Mensagem:
                    {message}
                    """
                )
            except Exception as e:
                print(f'Erro ao enviar email para admin {admin_email}: {str(e)}')
        
        return jsonify({'success': True})
    except Exception as e:
        return jsonify({'success': False, 'error': str(e)}), 500

@app.route('/status')
@login_required
def status():
    user_id = get_user_id()
    plan, _, _, _ = get_user_plan_info(user_id)
    plan_data = PLANS.get(plan, PLANS['free'])
    
    # Obter status atual dos componentes
    api_status = check_api_status()
    db_status = check_database_status()
    auth_status = check_auth_status()
    payment_status = check_payment_status()
    
    # Obter incidentes recentes
    incidents = get_recent_incidents()
    
    return render_template('status.html',
        last_update=datetime.now().strftime('%d/%m/%Y %H:%M:%S'),
        api_status=api_status,
        db_status=db_status,
        auth_status=auth_status,
        payment_status=payment_status,
        incidents=incidents,
        plan_data=plan_data
    )

@app.route('/api/status')
def api_status():
    return jsonify({
        'api_status': check_api_status(),
        'db_status': check_database_status(),
        'auth_status': check_auth_status(),
        'payment_status': check_payment_status()
    })

def check_api_status():
    try:
        # Verificar se a API está respondendo
        response = requests.get('https://api.eazybrief.com/health', timeout=5)
        if response.status_code == 200:
            return 'operational'
        return 'degraded'
    except:
        return 'down'

def check_database_status():
    try:
        # Verificar conexão com o banco de dados
        db = firebase_admin.db.reference('/')
        db.get()
        return 'operational'
    except:
        return 'down'

def check_auth_status():
    try:
        # Verificar serviço de autenticação
        auth = firebase_admin.auth.Client()
        auth.list_users(max_results=1)
        return 'operational'
    except:
        return 'down'

def check_payment_status():
    try:
        # Verificar serviço de pagamento
        response = requests.get('https://api.stripe.com/v1/health', timeout=5)
        if response.status_code == 200:
            return 'operational'
        return 'degraded'
    except:
        return 'down'

def get_recent_incidents():
    try:
        # Buscar incidentes recentes do banco de dados
        incidents_ref = firebase_admin.db.reference('/incidents')
        incidents = incidents_ref.order_by_child('created_at').limit_to_last(5).get()
        
        if not incidents:
            return []
            
        # Converter para lista e ordenar por data
        incident_list = []
        for incident_id, incident in incidents.items():
            incident['id'] = incident_id
            incident_list.append(incident)
            
        return sorted(incident_list, key=lambda x: x['created_at'], reverse=True)
    except:
        return []

@app.route('/legal')
def legal():
    return render_template('legal.html',
        last_update=datetime.now().strftime('%d/%m/%Y')
    )

@app.route('/new_brief')
@login_required
def new_brief():
    user_id = get_user_id()
    plan, _, _, user_data = get_user_plan_info(user_id)
    return render_template('new_brief.html', plan_data=PLANS.get(plan, PLANS['free']))

@app.route('/feedback')
def feedback():
    return render_template('feedback.html')

@app.route('/api/feedback', methods=['POST'])
def submit_feedback():
    if not request.is_json:
        return jsonify({'success': False, 'error': 'Content-Type deve ser application/json'}), 400
    
    data = request.get_json()
    
    # Validar campos obrigatórios
    required_fields = ['type', 'title', 'description', 'priority']
    for field in required_fields:
        if not data.get(field):
            return jsonify({'success': False, 'error': f'Campo {field} é obrigatório'}), 400
    
    try:
        # Criar entrada no banco de dados
        feedback_data = {
            'type': data['type'],
            'title': data['title'],
            'description': data['description'],
            'priority': data['priority'],
            'contact': data.get('contact', ''),
            'status': 'pending',
            'created_at': datetime.utcnow().isoformat() + 'Z',
            'user_id': get_user_id() if 'user' in session else None,
            'user_email': session.get('user_email')
        }
        
        # Salvar no Firebase
        db.reference('feedback').push(feedback_data)
        
        # Enviar email para os admins
        admin_emails = ADMINS
        for admin_email in admin_emails:
            try:
                send_email(
                    subject=f'Novo Feedback: {data["title"]}',
                    recipients=[admin_email],
                    body=f"""
                    Novo feedback recebido:
                    
                    Tipo: {data['type']}
                    Título: {data['title']}
                    Descrição: {data['description']}
                    Prioridade: {data['priority']}
                    Contato: {data.get('contact', 'Não fornecido')}
                    Usuário: {session.get('user_email', 'Não logado')}
                    Data: {feedback_data['created_at']}
                    """
                )
            except Exception as e:
                print(f'Erro ao enviar email para admin {admin_email}: {str(e)}')
        
        return jsonify({'success': True})
    except Exception as e:
        print(f'Erro ao processar feedback: {str(e)}')
        return jsonify({'success': False, 'error': 'Erro ao processar feedback'}), 500

@app.context_processor
def inject_firebase_config():
    return {
        'config': {
            'FIREBASE_API_KEY': os.getenv('FIREBASE_API_KEY'),
            'FIREBASE_AUTH_DOMAIN': os.getenv('FIREBASE_AUTH_DOMAIN'),
            'FIREBASE_PROJECT_ID': os.getenv('FIREBASE_PROJECT_ID'),
            'FIREBASE_STORAGE_BUCKET': os.getenv('FIREBASE_STORAGE_BUCKET'),
            'FIREBASE_MESSAGING_SENDER_ID': os.getenv('FIREBASE_MESSAGING_SENDER_ID'),
            'FIREBASE_APP_ID': os.getenv('FIREBASE_APP_ID')
        }
    }

@login_required
@app.route('/export')
@login_required
def export_page():
    user_id = get_user_id()
    plan, _, _, user_data = get_user_plan_info(user_id)
    return render_template('export.html', plan_data=PLANS.get(plan, PLANS['free']))

@app.route('/api/export/<format>', methods=['POST'])
@login_required
def export_brief(format):
    user_id = get_user_id()
    plan, _, _, _ = get_user_plan_info(user_id)
    
    # Verifica se o formato está permitido no plano do usuário
    allowed_formats = PLANS[plan]['export_formats']
    if format not in allowed_formats:
        return jsonify({
            'error': f'Formato não disponível no seu plano {PLANS[plan]["name"]}. '
                    f'Formatos disponíveis: {", ".join(allowed_formats)}'
        }), 403
    
    if format not in ['pdf', 'docx', 'txt', 'html']:
        return jsonify({'error': 'Formato não suportado'}), 400
    
    data = request.get_json()
    options = {
        'includeMetadata': data.get('includeMetadata', False),
        'includeImages': data.get('includeImages', False),
        'includeComments': data.get('includeComments', False)
    }
    
    try:
        if format == 'pdf':
            return export_brief_pdf(options)
        elif format == 'docx':
            return export_brief_docx(options)
        elif format == 'txt':
            return export_brief_txt(options)
        elif format == 'html':
            return export_brief_html(options)
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Rotas de integração
@app.route('/api/integrations/google-drive/auth')
@login_required
def google_drive_auth():
    try:
        from google_auth_oauthlib.flow import Flow
        from google.oauth2.credentials import Credentials
        from googleapiclient.discovery import build
        
        flow = Flow.from_client_secrets_file(
            'client_secrets.json',
            scopes=['https://www.googleapis.com/auth/drive.file']
        )
        flow.redirect_uri = url_for('google_drive_callback', _external=True)
        
        authorization_url, state = flow.authorization_url(
            access_type='offline',
            include_granted_scopes='true'
        )
        
        return jsonify({'authUrl': authorization_url})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/integrations/trello/auth')
@login_required
def trello_auth():
    try:
        from trello import TrelloClient
        
        client = TrelloClient(
            api_key=os.getenv('TRELLO_API_KEY'),
            api_secret=os.getenv('TRELLO_API_SECRET')
        )
        
        auth_url = client.get_authorization_url(
            name='EazyBrief',
            expiration='never',
            scope=['read', 'write']
        )
        
        return jsonify({'authUrl': auth_url})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/integrations/slack/auth')
@login_required
def slack_auth():
    try:
        client_id = os.getenv('SLACK_CLIENT_ID')
        scope = 'chat:write,channels:read,groups:read,im:read,mpim:read'
        redirect_uri = url_for('slack_callback', _external=True)
        
        auth_url = f'https://slack.com/oauth/v2/authorize?client_id={client_id}&scope={scope}&redirect_uri={redirect_uri}'
        
        return jsonify({'authUrl': auth_url})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

@app.route('/api/integrations/teams/auth')
@login_required
def teams_auth():
    try:
        client_id = os.getenv('TEAMS_CLIENT_ID')
        redirect_uri = url_for('teams_callback', _external=True)
        
        auth_url = f'https://login.microsoftonline.com/common/oauth2/v2.0/authorize?client_id={client_id}&response_type=code&redirect_uri={redirect_uri}&scope=offline_access%20Files.ReadWrite.All'
        
        return jsonify({'authUrl': auth_url})
    except Exception as e:
        return jsonify({'error': str(e)}), 500

# Callbacks de autenticação
@app.route('/integrations/google-drive/callback')
@login_required
def google_drive_callback():
    try:
        from google_auth_oauthlib.flow import Flow
        from google.oauth2.credentials import Credentials
        
        flow = Flow.from_client_secrets_file(
            'client_secrets.json',
            scopes=['https://www.googleapis.com/auth/drive.file']
        )
        flow.redirect_uri = url_for('google_drive_callback', _external=True)
        
        flow.fetch_token(code=request.args.get('code'))
        credentials = flow.credentials
        
        # Salvar credenciais no Firebase
        user_id = get_user_id()
        db.reference(f'user_integrations/{user_id}/google_drive').set({
            'token': credentials.token,
            'refresh_token': credentials.refresh_token,
            'token_uri': credentials.token_uri,
            'client_id': credentials.client_id,
            'client_secret': credentials.client_secret,
            'scopes': credentials.scopes
        })
        
        flash('Google Drive conectado com sucesso!', 'success')
        return redirect(url_for('export_page'))
    except Exception as e:
        flash(f'Erro ao conectar com Google Drive: {str(e)}', 'error')
        return redirect(url_for('export_page'))

@app.route('/integrations/trello/callback')
@login_required
def trello_callback():
    try:
        token = request.args.get('token')
        user_id = get_user_id()
        
        # Salvar token no Firebase
        db.reference(f'user_integrations/{user_id}/trello').set({
            'token': token
        })
        
        flash('Trello conectado com sucesso!', 'success')
        return redirect(url_for('export_page'))
    except Exception as e:
        flash(f'Erro ao conectar com Trello: {str(e)}', 'error')
        return redirect(url_for('export_page'))

@app.route('/integrations/slack/callback')
@login_required
def slack_callback():
    try:
        code = request.args.get('code')
        client_id = os.getenv('SLACK_CLIENT_ID')
        client_secret = os.getenv('SLACK_CLIENT_SECRET')
        
        # Trocar código por token
        response = requests.post('https://slack.com/api/oauth.v2.access', {
            'client_id': client_id,
            'client_secret': client_secret,
            'code': code
        })
        
        if response.ok:
            data = response.json()
            if data['ok']:
                user_id = get_user_id()
                db.reference(f'user_integrations/{user_id}/slack').set({
                    'access_token': data['access_token'],
                    'bot_user_id': data['bot_user_id'],
                    'team_id': data['team']['id'],
                    'team_name': data['team']['name']
                })
                
                flash('Slack conectado com sucesso!', 'success')
                return redirect(url_for('export_page'))
        
        raise Exception('Erro ao obter token do Slack')
    except Exception as e:
        flash(f'Erro ao conectar com Slack: {str(e)}', 'error')
        return redirect(url_for('export_page'))

@app.route('/integrations/teams/callback')
@login_required
def teams_callback():
    try:
        code = request.args.get('code')
        client_id = os.getenv('TEAMS_CLIENT_ID')
        client_secret = os.getenv('TEAMS_CLIENT_SECRET')
        redirect_uri = url_for('teams_callback', _external=True)
        
        # Trocar código por token
        response = requests.post('https://login.microsoftonline.com/common/oauth2/v2.0/token', {
            'client_id': client_id,
            'client_secret': client_secret,
            'code': code,
            'redirect_uri': redirect_uri,
            'grant_type': 'authorization_code'
        })
        
        if response.ok:
            data = response.json()
            user_id = get_user_id()
            db.reference(f'user_integrations/{user_id}/teams').set({
                'access_token': data['access_token'],
                'refresh_token': data['refresh_token'],
                'expires_in': data['expires_in']
            })
            
            flash('Microsoft Teams conectado com sucesso!', 'success')
            return redirect(url_for('export_page'))
        
        raise Exception('Erro ao obter token do Teams')
    except Exception as e:
        flash(f'Erro ao conectar com Microsoft Teams: {str(e)}', 'error')
        return redirect(url_for('export_page'))

def export_brief_pdf(options):
    try:
        from reportlab.lib import colors
        from reportlab.lib.pagesizes import letter
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        import io
        
        # Obter dados do brief
        brief_id = request.args.get('brief_id')
        brief_ref = db.reference(f'briefs/{brief_id}')
        brief_data = brief_ref.get()
        
        if not brief_data:
            raise Exception('Brief não encontrado')
        
        # Criar buffer para o PDF
        buffer = io.BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Título
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            spaceAfter=30
        )
        story.append(Paragraph(brief_data['title'], title_style))
        
        # Metadados (se solicitado)
        if options['includeMetadata']:
            meta_style = ParagraphStyle(
                'MetaData',
                parent=styles['Normal'],
                fontSize=10,
                textColor=colors.gray
            )
            story.append(Paragraph(f'Criado em: {brief_data.get("created_at", "N/A")}', meta_style))
            story.append(Paragraph(f'Última atualização: {brief_data.get("updated_at", "N/A")}', meta_style))
            story.append(Spacer(1, 20))
        
        # Conteúdo
        content_style = ParagraphStyle(
            'Content',
            parent=styles['Normal'],
            fontSize=12,
            spaceAfter=12
        )
        
        for section in brief_data.get('sections', []):
            story.append(Paragraph(section['title'], styles['Heading2']))
            story.append(Paragraph(section['content'], content_style))
            
            # Imagens (se solicitado)
            if options['includeImages'] and 'images' in section:
                for image_url in section['images']:
                    try:
                        img = Image(image_url, width=6*inch, height=4*inch)
                        story.append(img)
                        story.append(Spacer(1, 12))
                    except:
                        pass
            
            story.append(Spacer(1, 20))
        
        # Comentários (se solicitado)
        if options['includeComments'] and 'comments' in brief_data:
            story.append(Paragraph('Comentários', styles['Heading2']))
            for comment in brief_data['comments']:
                story.append(Paragraph(f"{comment['user']} ({comment['date']}):", styles['Heading3']))
                story.append(Paragraph(comment['text'], content_style))
                story.append(Spacer(1, 12))
        
        # Gerar PDF
        doc.build(story)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{brief_data['title']}.pdf",
            mimetype='application/pdf'
        )
    except Exception as e:
        raise Exception(f'Erro ao gerar PDF: {str(e)}')

def export_brief_docx(options):
    try:
        from docx import Document
        from docx.shared import Inches
        import io
        
        # Obter dados do brief
        brief_id = request.args.get('brief_id')
        brief_ref = db.reference(f'briefs/{brief_id}')
        brief_data = brief_ref.get()
        
        if not brief_data:
            raise Exception('Brief não encontrado')
        
        # Criar documento
        doc = Document()
        
        # Título
        doc.add_heading(brief_data['title'], 0)
        
        # Metadados (se solicitado)
        if options['includeMetadata']:
            doc.add_paragraph(f'Criado em: {brief_data.get("created_at", "N/A")}')
            doc.add_paragraph(f'Última atualização: {brief_data.get("updated_at", "N/A")}')
            doc.add_paragraph()
        
        # Conteúdo
        for section in brief_data.get('sections', []):
            doc.add_heading(section['title'], level=1)
            doc.add_paragraph(section['content'])
            
            # Imagens (se solicitado)
            if options['includeImages'] and 'images' in section:
                for image_url in section['images']:
                    try:
                        doc.add_picture(image_url, width=Inches(6))
                        doc.add_paragraph()
                    except:
                        pass
        
        # Comentários (se solicitado)
        if options['includeComments'] and 'comments' in brief_data:
            doc.add_heading('Comentários', level=1)
            for comment in brief_data['comments']:
                doc.add_heading(f"{comment['user']} ({comment['date']})", level=2)
                doc.add_paragraph(comment['text'])
        
        # Salvar em buffer
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        return send_file(
            buffer,
            as_attachment=True,
            download_name=f"{brief_data['title']}.docx",
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
    except Exception as e:
        raise Exception(f'Erro ao gerar DOCX: {str(e)}')

def export_brief_txt(options):
    try:
        # Obter dados do brief
        brief_id = request.args.get('brief_id')
        brief_ref = db.reference(f'briefs/{brief_id}')
        brief_data = brief_ref.get()
        
        if not brief_data:
            raise Exception('Brief não encontrado')
        
        # Criar conteúdo
        content = []
        content.append(f"Título: {brief_data['title']}\n")
        
        # Metadados (se solicitado)
        if options['includeMetadata']:
            content.append(f"Criado em: {brief_data.get('created_at', 'N/A')}")
            content.append(f"Última atualização: {brief_data.get('updated_at', 'N/A')}\n")
        
        # Conteúdo
        for section in brief_data.get('sections', []):
            content.append(f"\n{section['title']}")
            content.append("=" * len(section['title']))
            content.append(f"\n{section['content']}\n")
            
            # URLs das imagens (se solicitado)
            if options['includeImages'] and 'images' in section:
                content.append("Imagens:")
                for image_url in section['images']:
                    content.append(f"- {image_url}")
                content.append("")
        
        # Comentários (se solicitado)
        if options['includeComments'] and 'comments' in brief_data:
            content.append("\nComentários")
            content.append("=" * 10)
            for comment in brief_data['comments']:
                content.append(f"\n{comment['user']} ({comment['date']}):")
                content.append(f"{comment['text']}\n")
        
        # Gerar arquivo
        output = "\n".join(content)
        
        return Response(
            output,
            mimetype='text/plain',
            headers={
                'Content-Disposition': f'attachment; filename={brief_data["title"]}.txt'
            }
        )
    except Exception as e:
        raise Exception(f'Erro ao gerar TXT: {str(e)}')


    
    return Response(
        html_content,
        mimetype='text/html',
        headers={
            'Content-Disposition': f'attachment; filename="brief_{brief_id}.html"'
        }
    )

@app.context_processor
def inject_user_theme():
    user_theme = 'light'
    user_photo = None
    user_plan = 'free'
    user_id = get_user_id()
    if user_id:
        prefs = db.reference(f'user_prefs/{user_id}').get() or {}
        user_theme = prefs.get('theme', 'light')
        user_info = auth.get_user_by_email(session['user_email'])
        user_photo = prefs.get('photo_url', user_info.photo_url)
        
        # Buscar plano diretamente do banco de dados
        user_data = db.reference(f'users/{user_id}').get() or {}
        user_plan = user_data.get('plan', 'free')
        
        # Log para debug
        logging.info(f'[Context] Injetando contexto para user_id={user_id}')
        logging.info(f'[Context] Dados do usuário: {user_data}')
        logging.info(f'[Context] Plano atual: {user_plan}')
    
    return dict(user_theme=user_theme, user_photo=user_photo, user_plan=user_plan, plans=PLANS)

def get_oauth2_config():
    return {
        "web": {
            "client_id": os.getenv('GOOGLE_CLIENT_ID'),
            "client_secret": os.getenv('GOOGLE_CLIENT_SECRET'),
            "auth_uri": "https://accounts.google.com/o/oauth2/auth",
            "token_uri": "https://oauth2.googleapis.com/token",
            "redirect_uris": [f"{request.url_root}oauth2callback"]
        }
    }

@app.route('/authorize')
def authorize():
    flow = Flow.from_client_config(
        get_oauth2_config(),
        scopes=['https://www.googleapis.com/auth/calendar'],
        redirect_uri=f"{request.url_root}oauth2callback"
    )
    authorization_url, state = flow.authorization_url(
        access_type='offline',
        include_granted_scopes='true'
    )
    session['state'] = state
    return redirect(authorization_url)

@app.route('/oauth2callback')
def oauth2callback():
    state = session['state']
    flow = Flow.from_client_config(
        get_oauth2_config(),
        scopes=['https://www.googleapis.com/auth/calendar'],
        redirect_uri=f"{request.url_root}oauth2callback",
        state=state
    )
    flow.fetch_token(
        authorization_response=request.url,
        code=request.args.get('code')
    )
    credentials = flow.credentials
    session['credentials'] = {
        'token': credentials.token,
        'refresh_token': credentials.refresh_token,
        'token_uri': credentials.token_uri,
        'client_id': credentials.client_id,
        'client_secret': credentials.client_secret,
        'scopes': credentials.scopes
    }
    return redirect(url_for('index'))

@app.route('/api/transcribe-audio', methods=['POST'])
def transcribe_audio():
    if 'audio' not in request.files:
        return jsonify({'error': 'Nenhum arquivo de áudio enviado'}), 400

    audio_file = request.files['audio']
    if not audio_file.filename:
        return jsonify({'error': 'Nome do arquivo inválido'}), 400

    try:
        # Salvar o arquivo temporariamente
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, secure_filename(audio_file.filename))
        audio_file.save(temp_path)

        # Transcrever usando a API do Groq
        with open(temp_path, "rb") as file:
            transcription = groq_client.audio.transcriptions.create(
                file=(temp_path, file.read()),
                model="whisper-large-v3-turbo",
                response_format="verbose_json"
            )

        # Limpar o arquivo temporário
        os.remove(temp_path)

        return jsonify({
            'text': transcription.text,
            'success': True
        })

    except Exception as e:
        print(f"Erro ao transcrever áudio: {str(e)}")
        if os.path.exists(temp_path):
            os.remove(temp_path)
        return jsonify({'error': f'Erro ao transcrever áudio: {str(e)}'}), 500

if __name__ == '__main__':
    app.run(debug=True, port=5000)